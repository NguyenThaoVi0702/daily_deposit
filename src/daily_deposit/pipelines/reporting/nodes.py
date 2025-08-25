import pandas as pd
import numpy as np
import lightgbm as lgb
from sklearn.metrics import mean_absolute_error, median_absolute_error
import matplotlib.pyplot as plt
from typing import Dict, Any, List
import io
import logging

import io
from docx import Document 
from docx.shared import Inches

log = logging.getLogger(__name__)


from daily_deposit.pipelines.data_engineering.nodes import create_features_for_model1, create_intraday_features_for_model2


def _quantile_loss(y_true, y_pred, alpha):
    error = y_true - y_pred
    return np.mean(np.maximum(alpha * error, (alpha - 1) * error))



def get_test_set(
    daily_data: pd.DataFrame,
    hourly_data: pd.DataFrame,
    m1_training_preds: pd.DataFrame,
) -> (pd.DataFrame, pd.DataFrame, pd.Series):
    """Determines the test set based on the last date from M1's training predictions."""
    daily_data['cob_dt'] = pd.to_datetime(daily_data['cob_dt'])
    hourly_data['cob_dt'] = pd.to_datetime(hourly_data['cob_dt'])
    m1_training_preds['cob_dt'] = pd.to_datetime(m1_training_preds['cob_dt'])

    last_training_date = m1_training_preds['cob_dt'].max()
    
    df_daily_test = daily_data[daily_data['cob_dt'] > last_training_date].copy()
    df_hourly_test = hourly_data[hourly_data['cob_dt'] > last_training_date].copy()
    test_dates = pd.Series(df_daily_test['cob_dt'].unique()).sort_values()
    
    return df_daily_test, df_hourly_test, test_dates


def generate_test_predictions(
    test_dates: pd.Series,
    daily_data_full: pd.DataFrame,
    hourly_data_test: pd.DataFrame,
    params_m1_feature_creation: Dict, 
    params_reporting: Dict,
    model_m1_q01, model_m1_q05, model_m1_q09, model_m1_mean,
    model_m2_median, model_m2_mean,
) -> pd.DataFrame:
    """Generates all M1 and M2 predictions on the test set."""
    # Generate M1 Features
    lookback = pd.Timedelta(days=params_reporting["m1_feature_lookback_days"])
    
    log.info("--- [DEBUG] Slicing data for M1 Test Set Feature Creation ---")
    log.info(f"Test set starts on date: {test_dates.min().date()}")
    log.info(f"Lookback parameter is: {params_reporting['m1_feature_lookback_days']} days")
    log.info(f"Attempting to slice historical data from date: {(test_dates.min() - lookback).date()}")

    
    df_for_m1_features = daily_data_full[daily_data_full['cob_dt'] >= (test_dates.min() - lookback)].copy()


    log.info(f"Shape of the resulting slice `df_for_m1_features`: {df_for_m1_features.shape}")
    if not df_for_m1_features.empty:
        log.info(f"Actual date range in slice: {df_for_m1_features['cob_dt'].min().date()} to {df_for_m1_features['cob_dt'].max().date()}")


    df_m1_features = create_features_for_model1(df_for_m1_features, params_m1_feature_creation)


    if df_m1_features is None or df_m1_features.empty:
        raise ValueError(
            "Feature creation for M1 on the test set resulted in an empty DataFrame. "
            "This is likely due to an insufficient lookback period. "
            "Please check `m1_feature_lookback_days` in parameters or provide more historical data."
        )

    df_m1_features_test = df_m1_features.loc[df_m1_features.index.isin(test_dates)]

    # Predict with M1 Models
    m1_preds = {
        'm1_pred_q0.1': model_m1_q01.predict(df_m1_features_test[model_m1_q01.feature_name_]),
        'm1_pred_q0.5': model_m1_q05.predict(df_m1_features_test[model_m1_q05.feature_name_]),
        'm1_pred_q0.9': model_m1_q09.predict(df_m1_features_test[model_m1_q09.feature_name_]),
        'm1_pred_mean': model_m1_mean.predict(df_m1_features_test[model_m1_mean.feature_name_]),
    }
    df_m1_preds_test = pd.DataFrame(m1_preds, index=df_m1_features_test.index)

    # Generate M2 Features
    df_m2_features_test = create_intraday_features_for_model2(hourly_data_test, {"date_column": "cob_dt"})
    
    # Predict with M2 Models
    X_m2_base = pd.DataFrame(index=test_dates)
    X_m2_base['m1_pred_interval_width'] = df_m1_preds_test['m1_pred_q0.9'] - df_m1_preds_test['m1_pred_q0.1']
    X_m2_test_features = X_m2_base.join(df_m2_features_test, how='left').fillna(0)
    
    m2_error_preds = {
        'm2_median_error_pred': model_m2_median.predict(X_m2_test_features[model_m2_median.feature_name_]),
        'm2_mean_error_pred': model_m2_mean.predict(X_m2_test_features[model_m2_mean.feature_name_]),
    }
    df_m2_error_preds_test = pd.DataFrame(m2_error_preds, index=X_m2_test_features.index)

    return df_m1_preds_test.join(df_m2_error_preds_test, how='inner')


def generate_final_holdout_predictions(
    m1_features_test: pd.DataFrame,
    m2_features_test: pd.DataFrame,
    model_m1_q01, model_m1_q05, model_m1_q09, model_m1_mean,
    model_m2_median, model_m2_mean
) -> pd.DataFrame:
    """Generates all M1 and M2 predictions on the final holdout test set."""
    log.info("--- Generating predictions on final holdout test set ---")
    m1_preds = {
        'm1_pred_q0.1': model_m1_q01.predict(m1_features_test[model_m1_q01.feature_name_]),
        'm1_pred_q0.5': model_m1_q05.predict(m1_features_test[model_m1_q05.feature_name_]),
        'm1_pred_q0.9': model_m1_q09.predict(m1_features_test[model_m1_q09.feature_name_]),
        'm1_pred_mean': model_m1_mean.predict(m1_features_test[model_m1_mean.feature_name_]),
    }
    df_m1_preds_test = pd.DataFrame(m1_preds, index=m1_features_test.index)

    X_m2_base = pd.DataFrame(index=m2_features_test.index)
    X_m2_base['m1_pred_interval_width'] = df_m1_preds_test['m1_pred_q0.9'] - df_m1_preds_test['m1_pred_q0.1']
    X_m2_test_features = X_m2_base.join(m2_features_test, how='left').fillna(0)

    m2_error_preds = {
        'm2_median_error_pred': model_m2_median.predict(X_m2_test_features[model_m2_median.feature_name_]),
        'm2_mean_error_pred': model_m2_mean.predict(X_m2_test_features[model_m2_mean.feature_name_])
    }
    df_m2_error_preds_test = pd.DataFrame(m2_error_preds, index=X_m2_test_features.index)
    
    return df_m1_preds_test.join(df_m2_error_preds_test, how='inner')

def create_evaluation_table(
    predictions: pd.DataFrame,
    daily_test_set: pd.DataFrame,
    params_m1_training: Dict,
    params_reporting: Dict,
) -> pd.DataFrame:
    """Combines predictions with actuals and calculates final corrected forecasts."""
    results = predictions.copy()
    

    actuals = daily_test_set[params_m1_training["target_col"]].rename('actual')
    results = results.join(actuals).fillna(0)

    factors = params_reporting["m2_correction_factors"]
    results['m2_dampened_mean_adjustment'] = results['m2_mean_error_pred'] * factors['mean']
    results['final_pred_mean'] = results['m1_pred_mean'] + results['m2_dampened_mean_adjustment']
    results['m2_dampened_median_adjustment'] = results['m2_median_error_pred'] * factors['median']
    for alpha in [0.1, 0.5, 0.9]:
        results[f'final_pred_q{alpha}'] = results[f'm1_pred_q{alpha}'] + results['m2_dampened_median_adjustment']
    return results
    

def analyze_point_forecasts(results_df: pd.DataFrame) -> Dict:
    """Analyzes MAE and MedAE for mean and median forecasts."""
    metrics = {}
    for forecast_type in ['median', 'mean']:
        prefix = 'q0.5' if forecast_type == 'median' else 'mean'
        m1_col, final_col = f'm1_pred_{prefix}', f'final_pred_{prefix}'
        
        mae_m1 = mean_absolute_error(results_df['actual'], results_df[m1_col])
        mae_final = mean_absolute_error(results_df['actual'], results_df[final_col])
        improvement = (mae_m1 - mae_final) / mae_m1 if mae_m1 > 1e-9 else 0
        
        metrics[forecast_type] = {
            "MAE M1": mae_m1, "MAE Final": mae_final, "MAE Improvement": improvement,
            "MedAE M1": median_absolute_error(results_df['actual'], results_df[m1_col]),
            "MedAE Final": median_absolute_error(results_df['actual'], results_df[final_col]),
        }
    return metrics

def analyze_interval_forecasts(results_df: pd.DataFrame) -> Dict:
    """Analyzes coverage and sharpness of the 80% prediction interval."""
    m1_lower, m1_upper = results_df['m1_pred_q0.1'], results_df['m1_pred_q0.9']
    final_lower, final_upper = results_df['final_pred_q0.1'], results_df['final_pred_q0.9']
    
    return {
        "M1 Coverage": ((results_df['actual'] >= m1_lower) & (results_df['actual'] <= m1_upper)).mean(),
        "Final Coverage": ((results_df['actual'] >= final_lower) & (results_df['actual'] <= final_upper)).mean(),
        "M1 Width": (m1_upper - m1_lower).mean(),
        "Final Width": (final_upper - final_lower).mean(),
    }
    
def analyze_quantile_losses(results_df: pd.DataFrame) -> Dict:
    """Analyzes the pinball loss for each quantile."""
    losses = {}
    for alpha in [0.1, 0.5, 0.9]:
        loss_m1 = _quantile_loss(results_df['actual'], results_df[f'm1_pred_q{alpha}'], alpha)
        loss_final = _quantile_loss(results_df['actual'], results_df[f'final_pred_q{alpha}'], alpha)
        losses[f'q{alpha}'] = {"M1 Loss": loss_m1, "Final Loss": loss_final}
    return losses
    


def plot_feature_importance(model: lgb.LGBMRegressor, model_name: str) -> (plt.Figure, plt.Figure):
    """
    Creates a feature importance plot, formats the numbers, and returns the
    figure object twice for two different outputs (memory and file).
    """
    fig, ax = plt.subplots(figsize=(10, 8))
    lgb.plot_importance(
        model, 
        ax=ax, 
        max_num_features=20, 
        importance_type='gain', 
        title=f"Feature Importance (Gain) for {model_name}"
    )


    for text_object in ax.texts:
        try:

            current_value = float(text_object.get_text())

            text_object.set_text(f'{current_value:.2e}')
        except ValueError:

            pass


    plt.tight_layout()
    

    return fig, fig

def plot_forecast_comparison(results_df: pd.DataFrame, forecast_type: str, factors: Dict) -> plt.Figure:
    """Generates the main forecast comparison plot."""

    fig, ax = plt.subplots(figsize=(20, 9))
    ax.fill_between(results_df.index, results_df['final_pred_q0.1'], results_df['final_pred_q0.9'], color='skyblue', alpha=0.4, label='Final 80% Interval')
    ax.plot(results_df.index, results_df['actual'], 'o-', color='black', label='Actual')
    
    if forecast_type == 'median':
        final_col, title = 'final_pred_q0.5', 'Median'
    else:
        final_col, title = 'final_pred_mean', 'Mean'
        
    ax.plot(results_df.index, results_df[final_col], '--', color='red', label=f'Final {title} Prediction')
    ax.set_title(f'Final System Forecast vs. Actual ({title})')
    ax.legend()
    return fig

# --- The Final Report Generation Node ---


def generate_docx_report(
    point_metrics: Dict,
    interval_metrics: Dict,
    quantile_losses: Dict,
    fi_m1_q01: plt.Figure, fi_m1_q05: plt.Figure, fi_m1_q09: plt.Figure, fi_m1_mean: plt.Figure,
    fi_m2_median: plt.Figure, fi_m2_mean: plt.Figure,
    plot_comp_median: plt.Figure, plot_comp_mean: plt.Figure
) -> None:
    """Assembles key metrics, analysis, and plots into a final .docx report."""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        raise ImportError("Please install `python-docx` to generate reports.")

    document = Document()
    document.add_heading('Daily Loan Prediction System Evaluation Report', 0)
    document.add_paragraph(f"Report generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")

    # Helper to add a plot
    def add_plot(fig):
        memfile = io.BytesIO()
        fig.savefig(memfile, format='png', bbox_inches='tight')
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6.0))
        memfile.close()
        plt.close(fig)

    # --- Section 1: Point Forecast Performance ---
    document.add_heading('1. Point Forecast Performance (Mean Absolute Error)', level=1)
    table = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'Forecast Type', 'M1 MAE', 'Final MAE', 'Improvement'
    for key, value in point_metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key.capitalize()
        row_cells[1].text = f'{value["MAE M1"]:,.0f}'
        row_cells[2].text = f'{value["MAE Final"]:,.0f}'
        row_cells[3].text = f'{value["MAE Improvement"]:.2%}'
    document.add_paragraph() 

    # --- Section 2: Prediction Interval Performance ---
    document.add_heading('2. Prediction Interval Performance (80% Interval)', level=1)
    p_interval = document.add_paragraph()
    p_interval.add_run('M1 Original Interval:\t').bold = True
    p_interval.add_run(f"Coverage = {interval_metrics['M1 Coverage']:.2%},  Avg. Width = {interval_metrics['M1 Width']:,.0f}\n")
    p_interval.add_run('Final Shifted Interval:\t').bold = True
    p_interval.add_run(f"Coverage = {interval_metrics['Final Coverage']:.2%},  Avg. Width = {interval_metrics['Final Width']:,.0f}")
    document.add_paragraph()

    # --- Section 3: Quantile Loss Performance ---
    document.add_heading('3. Quantile (Pinball) Loss Performance', level=1)
    table_ql = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells_ql = table_ql.rows[0].cells
    hdr_cells_ql[0].text, hdr_cells_ql[1].text, hdr_cells_ql[2].text = 'Quantile', 'M1 Loss', 'Final Loss'
    for key, value in quantile_losses.items():
        row_cells_ql = table_ql.add_row().cells
        row_cells_ql[0].text = key
        row_cells_ql[1].text = f'{value["M1 Loss"]:,.0f}'
        row_cells_ql[2].text = f'{value["Final Loss"]:,.0f}'
    document.add_paragraph()

    # --- Section 4: Visualizations ---
    document.add_heading('4. Forecast Visualizations', level=1)
    document.add_paragraph('Comparison plot for the Median forecast:')
    add_plot(plot_comp_median)
    document.add_paragraph('Comparison plot for the Mean forecast:')
    add_plot(plot_comp_mean)
    document.add_page_break()

    # --- Section 5: Feature Importances ---
    document.add_heading('5. Feature Importances (Top 20 by Gain)', level=1)
    document.add_heading('Model 1 Importances', level=2)
    add_plot(fi_m1_mean)
    add_plot(fi_m1_q05)
    document.add_heading('Model 2 Corrector Importances', level=2)
    add_plot(fi_m2_mean)
    add_plot(fi_m2_median)
    
    # Save the document
    filepath = "data/08_reporting/final_summary_report.docx"
    document.save(filepath)
    log.info(f"Report successfully generated at {filepath}")
    
    
def generate_m1_only_holdout_predictions(
    m1_features_test: pd.DataFrame,
    model_m1_q01, model_m1_q05, model_m1_q09, model_m1_mean
) -> pd.DataFrame:
    """
    Generates M1 predictions on the final holdout test set.
    """
    log.info("--- Generating M1-only predictions on final holdout test set ---")
    m1_preds = {
        'm1_pred_q0.1': model_m1_q01.predict(m1_features_test[model_m1_q01.feature_name_]),
        'm1_pred_q0.5': model_m1_q05.predict(m1_features_test[model_m1_q05.feature_name_]),
        'm1_pred_q0.9': model_m1_q09.predict(m1_features_test[model_m1_q09.feature_name_]),
        'm1_pred_mean': model_m1_mean.predict(m1_features_test[model_m1_mean.feature_name_]),
    }
    df_m1_preds_test = pd.DataFrame(m1_preds, index=m1_features_test.index)
    return df_m1_preds_test


def create_m1_only_evaluation_table(
    m1_predictions: pd.DataFrame,
    daily_test_set: pd.DataFrame,
    params_m1_training: Dict,
) -> pd.DataFrame:
    """
    Combines M1-only predictions with actuals. The 'final' prediction is just the M1 prediction.
    """
    log.info("--- Creating M1-only evaluation table ---")
    results = m1_predictions.copy()
    actuals = daily_test_set[params_m1_training["target_col"]].rename('actual')
    results = results.join(actuals).fillna(0)


    results['final_pred_mean'] = results['m1_pred_mean']
    results['final_pred_q0.1'] = results['m1_pred_q0.1']
    results['final_pred_q0.5'] = results['m1_pred_q0.5']
    results['final_pred_q0.9'] = results['m1_pred_q0.9']
    
    return results


def plot_m1_only_forecast_comparison(results_df: pd.DataFrame, forecast_type: str) -> plt.Figure:
    """
    Generates a comparison plot showing only M1 predictions against actuals.
    """
    fig, ax = plt.subplots(figsize=(20, 9))
    ax.fill_between(results_df.index, results_df['m1_pred_q0.1'], results_df['m1_pred_q0.9'], color='skyblue', alpha=0.4, label='M1 80% Interval')
    ax.plot(results_df.index, results_df['actual'], 'o-', color='black', label='Actual Value')
    
    if forecast_type == 'median':
        m1_col, title, fcolor = 'm1_pred_q0.5', 'Median', 'red'
    else:
        m1_col, title, fcolor = 'm1_pred_mean', 'Mean', 'purple'
        
    ax.plot(results_df.index, results_df[m1_col], '--', color=fcolor, linewidth=2.5, label=f'M1 {title} Prediction')
    ax.set_title(f'Model 1 Forecast vs. Actual ({title}) on Holdout Set')
    ax.legend()
    return fig


def generate_m1_only_docx_report(
    point_metrics: Dict,
    interval_metrics: Dict,
    quantile_losses: Dict,
    fi_m1_q01: plt.Figure, fi_m1_q05: plt.Figure, fi_m1_q09: plt.Figure, fi_m1_mean: plt.Figure,
    plot_comp_median: plt.Figure, plot_comp_mean: plt.Figure
) -> None:
    """
    Assembles all M1-only metrics and plots into a single .docx report.
    """
    from docx import Document
    from docx.shared import Inches

    document = Document()
    document.add_heading('Model 1 Prediction System Evaluation Report', 0)
    document.add_paragraph(f"Report generated on: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    
    def add_plot(fig):
        memfile = io.BytesIO()
        fig.savefig(memfile, format='png', bbox_inches='tight')
        memfile.seek(0)
        document.add_picture(memfile, width=Inches(6.0))
        memfile.close()
        plt.close(fig)

    document.add_heading('1. Point Forecast Performance (MAE)', level=1)
    table = document.add_table(rows=1, cols=3, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'Forecast Type', 'M1 MAE', 'M1 MedAE'
    for key, value in point_metrics.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key.capitalize()
        row_cells[1].text = f'{value["MAE M1"]:,.0f}'
        row_cells[2].text = f'{value["MedAE M1"]:,.0f}'
    document.add_paragraph()
    
    document.add_heading('2. Prediction Interval Performance (80% Interval)', level=1)
    p_interval = document.add_paragraph()
    p_interval.add_run('M1 Interval:\t').bold = True
    p_interval.add_run(f"Coverage = {interval_metrics['M1 Coverage']:.2%},  Avg. Width = {interval_metrics['M1 Width']:,.0f}")
    document.add_paragraph()

    document.add_heading('3. Quantile (Pinball) Loss Performance', level=1)
    table_ql = document.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells_ql = table_ql.rows[0].cells
    hdr_cells_ql[0].text, hdr_cells_ql[1].text = 'Quantile', 'M1 Loss'
    for key, value in quantile_losses.items():
        row_cells_ql = table_ql.add_row().cells
        row_cells_ql[0].text = key
        row_cells_ql[1].text = f'{value["M1 Loss"]:,.0f}'
    document.add_paragraph()

    document.add_heading('4. Forecast Visualizations', level=1)
    add_plot(plot_comp_median)
    add_plot(plot_comp_mean)
    document.add_page_break()

    document.add_heading('5. Feature Importances (Top 20 by Gain)', level=1)
    add_plot(fi_m1_mean)
    add_plot(fi_m1_q05)
    
    filepath = "data/08_reporting/m1_only_summary_report.docx"
    document.save(filepath)
    log.info(f"M1-only report successfully generated at {filepath}")
